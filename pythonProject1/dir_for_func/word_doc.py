from docx.oxml.shared import OxmlElement, qn
import docx
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from num_to_rus import Converter
import numpy as np
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

def create_word_files(self,dfs):




    folder_path = os.path.join(os.getcwd(), 'addition')

    # Создаем новую папку в "addition"
    new_folder_path = os.path.join(folder_path, 'Протоколы')
    if not os.path.exists(new_folder_path):
        os.makedirs(new_folder_path)
        print('Папка "Протоколы" успешно создана.')
    else:
        print('Папка "Протоколы" уже существует, пропускаем ее создание.')

    # Создаем файл Word "ПРОТОКОЛ №1"
    document = Document()
    document.add_heading('ПРОТОКОЛ №1', 0)
    document.add_paragraph('Содержимое протокола...')
    file_path = os.path.join(new_folder_path, 'ПРОТОКОЛ №1.docx')
    document.save(file_path)

    try:
        # Извлекаем нужные DataFrame
        df = dfs['Конкурс (Работники)']
        df1 = dfs['Конкурс (Комиссия)']

        # Создаем списки для хранения данных
        name_workers = []

        # Сортируем DataFrame
        sorted_df = df.sort_values(by=['Новый отдел', 'Группа', 'Категория'])
        # Дотсаем все баллы по фамилиям кандидатов, чтоб потом их сформировать рейтинг по баллам
        result_names = {}
        for index, row in sorted_df.iterrows():
            fio = row['ФИО']
            result_names[fio] = [row[col] for col in sorted_df.columns if
                                col.endswith('Балл') or col.startswith('Баллы')]
        raiting= reiting_candidates(result_names)
        full_raiting_with_stages=first_table(raiting,file_path)
        full_new = {}
        unfull = {}
        for key, value in full_raiting_with_stages.items():
            if value in ('неявка', 'Неявка'):
                if key not in unfull:
                    unfull[key] = value
            else:
                if key not in full_new:
                    full_new[key] = value

        # Создаем словарь, где ключи - это ФИО из df1
        data = {}
        for col in df1.columns:
            if 'ФИО' in col:
                for i, value in enumerate(df1[col]):
                    if value not in data:
                        data[value] = {
                            'ball': [],
                            'za': [],
                            'motivation': []
                        }

        # Заполняем словарь данными из sorted_df
        for col_idx, col in enumerate(sorted_df.columns):
            col_words = col.split()[:3]
            col_words_str = ' '.join(col_words)
            for key in data:
                words = key.split()[:3]
                words_str = ' '.join(words)
                if words_str == col_words_str:
                    if col.endswith('Балл'):
                        data[key]['ball'].extend(sorted_df[col].tolist())
                    elif col.endswith('ЗаПротив'):
                        data[key]['za'].extend(sorted_df[col].tolist())
                    elif col.endswith('Мотивировка'):
                        data[key]['motivation'].extend(sorted_df[col].tolist())

        # Извлекаем ФИО из sorted_df
        for col_idx, col in enumerate(sorted_df.columns):
            if 'ФИО' in col:
                name_workers.extend(sorted_df[col].tolist())
        print(name_workers)
        for i, key in enumerate(name_workers):
                document = add_rating_table(file_path, key, data, i, full_new)
                document.save(file_path)
        print('Файл "ПРОТОКОЛ №1.docx" успешно создан!')
    except Exception as e:
        print(f'Ошибка при создании файла: {e}')

def add_rating_table(file_path, candidate_name, data, num, check ):
    # Открываем существующий документ Word
    conv = Converter()

    member_name = []
    for val in data:
        member_name.append(val)


    document = Document(file_path)
    if (candidate_name)  in check:
        desired_value = conv.convert(check[candidate_name])

        # Добавляем таблицу
        dlina=4+len(member_name)+1
        table = document.add_table(rows=dlina, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Устанавливаем стиль границ таблицы
        tbl = table._tbl

        # Устанавливаем шрифт Times New Roman для всей таблицы
        # Устанавливаем шрифт Times New Roman для всей таблицы
        tcFont = OxmlElement('w:tcPr')
        font = OxmlElement('w:rFonts')
        font.set(qn('w:ascii'), 'Times New Roman')
        font.set(qn('w:hAnsi'), 'Times New Roman')
        tcFont.append(font)
        for row in tbl.tr_lst:
            for cell in row.tc_lst:
                tcPr = cell.get_or_add_tcPr()
                tcPr.append(tcFont)

        # Устанавливаем стиль границ таблицы
        for row_idx, row in enumerate(tbl.tr_lst):
            for cell in row.tc_lst:
                tcPr = cell.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')



                # Только для ячеек, начиная со второй строки (row_idx != 1)
                if row_idx != 1:
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                    top.set(qn('w:sz'), '4')
                    top.set(qn('w:color'), '000000')
                    tcBorders.append(top)

                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), '000000')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '4')
                left.set(qn('w:color'), '000000')
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:sz'), '4')
                right.set(qn('w:color'), '000000')

                tcBorders.append(bottom)
                tcBorders.append(left)
                tcBorders.append(right)
                tcPr.append(tcBorders)



            # Заполняем первую и вторую ячейки (ФИО кандидата и его данные)
            cell = table.cell(0, 0)
            cell.merge(table.cell(0, 1))
            cell.merge(table.cell(0, 2))
            cell.merge(table.cell(0, 3))
            cell.text = candidate_name
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].underline = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            cell.width = Inches(8)

            cell = table.cell(1, 0)
            cell.merge(table.cell(1, 1))
            cell.merge(table.cell(1, 2))
            cell.merge(table.cell(1, 3))
            cell.text = f"Фамилия, имя, отчество кандидата, занявшего место № {desired_value}  в рейтинге"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.width = Inches(8)

            # Заполняем третью и четвертую ячейки (Голосование)

            cell = table.cell(2, 0)
            cell.merge(table.cell(3, 0))
            cell.text = "Фамилия, имя, отчество члена конкурсной комиссии"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.width = Inches(8)

            cell = table.cell(2, 1)
            cell.merge(table.cell(2, 2))
            cell.merge(table.cell(2, 3))
            cell.text = "Голосование"
            cell.paragraphs[0].runs[0].font.bold = False
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Inches(8)

            cell = table.cell(3, 1)
            cell.text = '"за"'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Inches(2.67)

            cell = table.cell(3, 2)
            cell.text = '"против"'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Inches(2.67)

            cell = table.cell(3, 3)
            cell.text = '"воздержался"'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Inches(2.66)

            # Заполняем пятую и шестую ячейки (Фамилия, имя, отчество члена конкурсной комиссии)
            if member_name:
                for i in range(len(member_name)):
                    cell = table.cell(i+4,0)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.text = f" {member_name[i]}"
            else:
                cell = table.cell(4, 0)
                cell.text = "Нет данных о членах конкурсной комиссии"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        reshenie=[]
        for name in member_name:
            if len(data[name]["za"]) > num:
                if str(data[name]["za"][num]) == 'nan':
                    reshenie.append('Воздержусь')
                else:
                    reshenie.append(data[name]["za"][num])

            else:
                reshenie.append('Воздержусь')


        # Заполняем ячейки таблицы в зависимости от значений в reshenie
        # Заполняем остальные ячейки
        o = 4

        for j, value in enumerate(reshenie):

            for i in range(o, dlina-1):

                if value == 'За':
                    table.cell(i, 1).text = '+'
                    table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
                elif value == 'Против':
                    table.cell(i, 2).text = '+'
                    table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
                elif value == 'Воздержусь':
                    table.cell(i, 3).text = '+'
                    table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
            o+=1

        # Заполняем последнюю строку с подведением итогов
        table.cell(dlina-1, 0).text = "Итого"
        table.cell(dlina-1, 0).paragraphs[0].runs[0].font.bold = True
        table.cell(dlina-1, 1).text = str(reshenie.count('За'))
        table.cell(dlina-1, 2).text = str(reshenie.count('Против'))
        table.cell(dlina-1, 3).text = str(reshenie.count('Воздержусь'))

        # Добавляем новый абзац с большим отступом после таблицы
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Inches(1)  # Устанавливаем отступ в 1 дюйм (2.54 см)
    # Сохраняем изменения в документе
    document.save(file_path)
    return document

def make_border(color, size):
    """Создает XML-элемент для границы ячейки"""
    border_element = OxmlElement('w:border')
    border_element.set(qn('w:color'), color.hex_l)
    border_element.set(qn('w:sz'), str(size.pt))
    border_element.set(qn('w:val'), 'single')
    return border_element

def reiting_candidates(data):

    raiting = {}
    for name, scores in data.items():
        valid_scores = []
        for score in scores:
            if isinstance(score, float) and np.isnan(score):
                continue
            elif score == 'Неявка':
                valid_scores.append('Неявка')
            else:
                try:
                    valid_scores.append(float(score))
                except ValueError:
                    continue
        if valid_scores:
            if 'Неявка' in valid_scores:
                raiting[name] = 'Неявка'
            else:
                avg_score = sum(valid_scores) / len(valid_scores)
                raiting[name] = avg_score
        else:
            raiting[name] = 0
    return raiting

def first_table(raiting, file_path):
    # Создаем новый документ Word
    doc = Document()

    # Создаем заголовок документа
    doc.add_heading("Рейтинг кандидатов", 0)

    # Сортируем словарь по значениям в обратном порядке, "Неявка" в конце
    sorted_raiting = sorted(raiting.items(), key=lambda x: (x[1] != 'Неявка', x[1] if x[1] != 'Неявка' else 0, x[0]), reverse=True)

    # Создаем таблицу
    table = doc.add_table(rows=1 + len(sorted_raiting), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = "ФИО"
    header_cells[1].text = "Балл"
    header_cells[2].text = "Место в рейтинге"

    # Создаем словарь для возврата
    result_raiting = {}

    # Заполняем таблицу и словарь
    current_place = 1
    for row_num, (name, score) in enumerate(sorted_raiting, start=1):
        row_cells = table.rows[row_num].cells
        if score == 'Неявка':
            row_cells[0].text = name
            row_cells[1].text = score
            row_cells[2].text = "Неявка"
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        else:
            row_cells[0].text = name
            row_cells[1].text = str(score)
            row_cells[2].text = str(current_place)
            result_raiting[name] = current_place
            current_place += 1

    # Добавляем отступ после таблицы
    doc.add_paragraph()

    # Сохраняем документ
    doc.save(file_path)

    return result_raiting